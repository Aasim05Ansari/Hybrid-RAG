from abc import ABC, abstractmethod
from html.parser import HTMLParser
from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader

from app.ingestion.metadata import Document


class DocumentLoader(ABC):

    @abstractmethod
    def load(self, path: Path) -> Document:
        """Load a file and return a Document."""
        raise NotImplementedError


class TextLoader(DocumentLoader):

    def load(self, path: Path) -> Document:

        if not path.exists():
            raise FileNotFoundError(
                f"File not found: {path}"
            )

        if not path.is_file():
            raise ValueError(
                f"Path is not a file: {path}"
            )

        content = path.read_text(
            encoding="utf-8"
        )

        return Document(
            content=content,
            metadata={
                "source": str(path),
                "filename": path.name,
                "file_type": path.suffix.lower(),
            },
        )


class HTMLTextExtractor(HTMLParser):

    def __init__(self):
        super().__init__()

        self.parts = []

        self.ignored_tags = {
            "script",
            "style",
            "noscript",
        }

        self.ignore_depth = 0

    def handle_starttag(self, tag, attrs):

        tag = tag.lower()

        if tag in self.ignored_tags:
            self.ignore_depth += 1
            return

        if self.ignore_depth > 0:
            return

        if tag in {
            "p",
            "div",
            "section",
            "article",
            "header",
            "footer",
            "main",
            "li",
            "h1",
            "h2",
            "h3",
            "h4",
            "h5",
            "h6",
            "br",
        }:
            self.parts.append("\n")

    def handle_endtag(self, tag):

        tag = tag.lower()

        if tag in self.ignored_tags:
            if self.ignore_depth > 0:
                self.ignore_depth -= 1
            return

        if self.ignore_depth > 0:
            return

        if tag in {
            "p",
            "div",
            "section",
            "article",
            "header",
            "footer",
            "main",
            "li",
            "h1",
            "h2",
            "h3",
            "h4",
            "h5",
            "h6",
        }:
            self.parts.append("\n")

    def handle_data(self, data):

        if self.ignore_depth > 0:
            return

        text = data.strip()

        if text:
            self.parts.append(text)

    def get_text(self) -> str:

        text = " ".join(self.parts)

        lines = [
            line.strip()
            for line in text.splitlines()
            if line.strip()
        ]

        return "\n\n".join(lines)


class HTMLLoader(DocumentLoader):

    def load(self, path: Path) -> Document:

        if not path.exists():
            raise FileNotFoundError(
                f"File not found: {path}"
            )

        if not path.is_file():
            raise ValueError(
                f"Path is not a file: {path}"
            )

        content = path.read_text(
            encoding="utf-8"
        )

        parser = HTMLTextExtractor()
        parser.feed(content)
        parser.close()

        extracted_text = parser.get_text()

        return Document(
            content=extracted_text,
            metadata={
                "source": str(path),
                "filename": path.name,
                "file_type": path.suffix.lower(),
            },
        )


class PDFLoader(DocumentLoader):

    def load(self, path: Path) -> Document:

        if not path.exists():
            raise FileNotFoundError(
                f"File not found: {path}"
            )

        if not path.is_file():
            raise ValueError(
                f"Path is not a file: {path}"
            )

        reader = PdfReader(str(path))

        pages = []

        for page in reader.pages:

            text = page.extract_text()

            if text:
                pages.append(text)

        content = "\n\n".join(pages)

        return Document(
            content=content,
            metadata={
                "source": str(path),
                "filename": path.name,
                "file_type": path.suffix.lower(),
                "page_count": len(reader.pages),
            },
        )


class DOCXLoader(DocumentLoader):

    def load(self, path: Path) -> Document:

        if not path.exists():
            raise FileNotFoundError(
                f"File not found: {path}"
            )

        if not path.is_file():
            raise ValueError(
                f"Path is not a file: {path}"
            )

        docx = DocxDocument(str(path))

        paragraphs = []

        for paragraph in docx.paragraphs:

            text = paragraph.text.strip()

            if text:
                paragraphs.append(text)

        content = "\n\n".join(paragraphs)

        return Document(
            content=content,
            metadata={
                "source": str(path),
                "filename": path.name,
                "file_type": path.suffix.lower(),
            },
        )


def get_loader(path: Path) -> DocumentLoader:

    extension = path.suffix.lower()

    if extension in {".txt", ".md", ".markdown"}:
        return TextLoader()

    if extension in {".html", ".htm"}:
        return HTMLLoader()

    if extension == ".pdf":
        return PDFLoader()

    if extension == ".docx":
        return DOCXLoader()

    raise ValueError(
        f"Unsupported file type: {extension}"
    )
