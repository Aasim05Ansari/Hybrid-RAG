from pathlib import Path

from app.ingestion.loader import (
    TextLoader,
    PDFLoader,
    DOCXLoader,
    get_loader,
)

def test_docx_loader(tmp_path):

    from docx import Document as DocxDocument

    file_path = tmp_path / "policy.docx"

    docx = DocxDocument()

    docx.add_paragraph(
        "Annual leave is 24 days per year."
    )

    docx.add_paragraph(
        "Sick leave is 12 days per year."
    )

    docx.save(file_path)

    loader = DOCXLoader()

    document = loader.load(file_path)

    assert (
        "Annual leave is 24 days per year."
        in document.content
    )

    assert (
        "Sick leave is 12 days per year."
        in document.content
    )

    assert document.metadata["filename"] == "policy.docx"

    assert document.metadata["file_type"] == ".docx"

def test_docx_loader_factory():

    loader = get_loader(
        Path("policy.docx")
    )

    assert isinstance(
        loader,
        DOCXLoader,
    )

def test_text_loader(tmp_path):

    file_path = tmp_path / "test.txt"

    file_path.write_text(
        "Annual leave is 24 days.",
        encoding="utf-8",
    )

    loader = TextLoader()

    document = loader.load(file_path)

    assert document.content == (
        "Annual leave is 24 days."
    )

    assert document.metadata["filename"] == "test.txt"

    assert document.metadata["file_type"] == ".txt"


def test_loader_factory():

    loader = get_loader(
        Path("policy.txt")
    )

    assert isinstance(
        loader,
        TextLoader,
    )
    
def test_pdf_loader(tmp_path):

    from pypdf import PdfWriter

    file_path = tmp_path / "policy.pdf"

    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)

    with open(file_path, "wb") as file:
        writer.write(file)

    loader = PDFLoader()

    document = loader.load(file_path)

    assert document.metadata["filename"] == "policy.pdf"
    assert document.metadata["file_type"] == ".pdf"
    assert document.metadata["page_count"] == 1